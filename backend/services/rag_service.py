"""
Service for Retrieval Augmented Generation (RAG)
"""
from typing import List, Dict, Any, Optional, Tuple
import httpx
import numpy as np
import os
import tempfile
import uuid
import logging
from sqlalchemy.orm import Session
from fastapi import UploadFile
import markdownify
import io
import PyPDF2
import docx
from datetime import datetime

from backend.config import settings
from backend.models.rag import RAGIndex, Document, DocumentChunk

# Configure logging
logger = logging.getLogger(__name__)

class RAGService:
    """Service for Retrieval Augmented Generation with database integration"""

    def __init__(self, db: Session):
        self.db = db
        self.embedding_url = "https://api.mistral.ai/v1/embeddings"
        self.api_key = settings.EXTERNAL_LLM_API_KEY

    async def create_index(self, name: str, user_id: Optional[int] = None,
                          team_id: Optional[int] = None,
                          deployed_app_id: Optional[int] = None,
                          embedding_model: str = "mistral-embed",
                          description: Optional[str] = None,
                          chunk_size: int = 1000,
                          chunk_overlap: int = 200) -> RAGIndex:
        """Create a new RAG index"""
        index = RAGIndex(
            name=name,
            description=description,
            user_id=user_id,
            team_id=team_id,
            deployed_app_id=deployed_app_id,
            embedding_model=embedding_model,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
        self.db.add(index)
        self.db.commit()
        self.db.refresh(index)
        return index

    async def upload_document(self, index_id: int, file: UploadFile) -> Document:
        """Upload a document to an index"""
        # Get the index
        index = self.db.query(RAGIndex).filter(RAGIndex.id == index_id).first()
        if not index:
            raise ValueError(f"Index with ID {index_id} not found")

        # Read file content
        file_content = await file.read()

        # Determine file type from filename
        file_extension = os.path.splitext(file.filename)[1].lower()
        file_type = file_extension.lstrip('.')

        # Create document record
        document = Document(
            index_id=index_id,
            filename=file.filename,
            file_type=file_type,
            file_size=len(file_content),
            content_type=file.content_type,
            external_id=str(uuid.uuid4())
        )

        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)

        # Process the document
        await self.process_document(document.id, file_content)

        return document

    async def process_document(self, document_id: int, file_content: bytes) -> None:
        """Process a document and create chunks with embeddings"""
        # Get the document
        document = self.db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise ValueError(f"Document with ID {document_id} not found")

        # Get the index
        index = document.index

        try:
            # Extract text from the document based on file type
            text = await self._extract_text(file_content, document.file_type)

            # Convert to markdown
            markdown_text = self._convert_to_markdown(text)

            # Store the full text content
            document.content = markdown_text

            # Split text into chunks
            chunks = self._split_text(markdown_text, index.chunk_size, index.chunk_overlap)

            # Generate embeddings for each chunk
            for i, chunk_text in enumerate(chunks):
                embedding = await self.generate_embedding(chunk_text)

                # Create chunk record
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    text=chunk_text,
                    embedding=embedding
                )

                self.db.add(chunk)

            # Mark document as processed
            document.is_processed = True

            self.db.commit()
            logger.info(f"Successfully processed document {document.id} with {len(chunks)} chunks")

        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}")
            # Update document status to indicate error
            document.is_processed = False
            document.doc_metadata = document.doc_metadata or {}
            document.doc_metadata["error"] = str(e)
            self.db.commit()
            raise

    async def generate_embedding(self, text: str) -> List[float]:
        """Generate embedding for text using Mistral's API"""
        async with httpx.AsyncClient() as client:
            response = await client.post(
                self.embedding_url,
                headers={"Authorization": f"Bearer {self.api_key}"},
                json={"model": "mistral-embed", "input": text}
            )

            if response.status_code != 200:
                raise Exception(f"Error generating embedding: {response.text}")

            result = response.json()
            return result["data"][0]["embedding"]

    async def search(self, index_id: int, query: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """Search for relevant chunks in an index"""
        # Get the index
        index = self.db.query(RAGIndex).filter(RAGIndex.id == index_id).first()
        if not index:
            raise ValueError(f"Index with ID {index_id} not found")

        # Generate embedding for the query
        query_embedding = await self.generate_embedding(query)

        # Get all chunks for the index
        chunks = self.db.query(DocumentChunk).join(Document).filter(Document.index_id == index_id).all()

        # Calculate similarity scores
        results = []
        for chunk in chunks:
            if not chunk.embedding:
                continue

            # Calculate cosine similarity
            similarity = self._cosine_similarity(query_embedding, chunk.embedding)

            results.append({
                "chunk_id": chunk.id,
                "document_id": chunk.document_id,
                "text": chunk.text,
                "similarity": similarity
            })

        # Sort by similarity and take top_k
        results.sort(key=lambda x: x["similarity"], reverse=True)
        return results[:top_k]

    async def augment_prompt(self, query: str, index_id: int, top_k: int = 3) -> str:
        """Augment the prompt with retrieved context"""
        # Search for relevant chunks
        results = await self.search(index_id, query, top_k)

        if not results:
            return query

        # Format the context
        context = "\n\n".join([f"Document chunk {i+1}:\n{result['text']}" for i, result in enumerate(results)])

        # Create the augmented prompt
        augmented_prompt = f"""Use the following information to answer the question. If the information provided doesn't contain the answer, just say that you don't know based on the provided information.

Context information:
{context}

Question: {query}

Answer:"""

        return augmented_prompt

    async def augment_messages(self, messages: List[Dict[str, str]], index_id: int, top_k: int = 3) -> List[Dict[str, str]]:
        """Augment chat messages with retrieved context"""
        if not messages:
            return messages

        # Get the last user message
        last_user_message = None
        for message in reversed(messages):
            if message.get("role") == "user":
                last_user_message = message
                break

        if not last_user_message:
            return messages

        # Search for relevant chunks
        results = await self.search(index_id, last_user_message["content"], top_k)

        if not results:
            return messages

        # Format the context
        context = "\n\n".join([f"Document chunk {i+1}:\n{result['text']}" for i, result in enumerate(results)])

        # Create a new system message with the context
        context_message = {
            "role": "system",
            "content": f"""Use the following information to answer the user's question. If the information provided doesn't contain the answer, just say that you don't know based on the provided information, and answer to the best of your general knowledge.

Context information:
{context}"""
        }

        # Find the system message if it exists
        system_message_index = None
        for i, message in enumerate(messages):
            if message.get("role") == "system":
                system_message_index = i
                break

        if system_message_index is not None:
            # Append the context to the existing system message
            messages[system_message_index]["content"] += f"\n\n{context_message['content']}"
        else:
            # Insert the context message at the beginning
            messages.insert(0, context_message)

        return messages

    def _cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)
        return np.dot(vec1, vec2) / (np.linalg.norm(vec1) * np.linalg.norm(vec2))

    async def _extract_text(self, file_content: bytes, file_type: str) -> str:
        """Extract text from a file based on its type"""
        if file_type in ['txt', 'md', 'markdown']:
            return file_content.decode('utf-8')

        elif file_type == 'pdf':
            return self._extract_text_from_pdf(file_content)

        elif file_type in ['docx', 'doc']:
            return self._extract_text_from_docx(file_content)

        else:
            # Default to treating as plain text
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from a PDF file"""
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        text = ""

        for page in reader.pages:
            text += page.extract_text() + "\n\n"

        return text

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from a DOCX file"""
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)

        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        return '\n'.join(full_text)

    def _convert_to_markdown(self, text: str) -> str:
        """Convert text to markdown format"""
        # Use markdownify to convert HTML to markdown if the text appears to be HTML
        if "<html" in text.lower() or "<body" in text.lower() or "<div" in text.lower():
            return markdownify.markdownify(text)

        # Otherwise, just return the text as is (it's already plain text or markdown)
        return text

    def _split_text(self, text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
        """Split text into chunks with overlap"""
        # Simple implementation - in production you'd want something more sophisticated
        words = text.split()
        chunks = []

        if not words:
            return []

        for i in range(0, len(words), chunk_size - chunk_overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)

        return chunks
